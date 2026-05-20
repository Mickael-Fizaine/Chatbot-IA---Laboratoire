"""
Chargeur de dataset large — PubMedQA pqa_artificial (211 000 articles)
=======================================================================
Charge 20 000 articles scientifiques biomedicaux et genere des fichiers
multi-formats dans data/raw/ pour tester le parser unstructured.io.

Fichiers generes dans data/raw/ :
  - 50 x .txt   (texte brut)
  - 30 x .docx  (rapport Word)
  - 15 x .pdf   (publication PDF)
  -  5 x .csv   (donnees experimentales tabulaires)

Usage :
  python data/load_large_dataset.py
  python data/load_large_dataset.py --docs 20000 --samples 100

Prérequis pip : datasets, fpdf2, python-docx
"""

import argparse
import csv
import json
import re
import sys
from pathlib import Path
from uuid import uuid4

ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(ROOT))

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from datasets import load_dataset
from tqdm import tqdm

RAW_DIR    = ROOT / "data" / "raw"
OUTPUT     = ROOT / "data" / "documents.json"
N_DOCS     = 20_000
N_SAMPLES  = 100   # files generated in data/raw/

# Distribution des formats de fichiers generes
FORMAT_SPLIT = {"txt": 50, "docx": 30, "pdf": 15, "csv": 5}


# ---------------------------------------------------------------------------
# Metadata extraction (reused from load_pubmedqa.py)
# ---------------------------------------------------------------------------

def _extract_compound(text: str) -> str:
    patterns = [
        r'\b([A-Z][a-z]+(?:inib|mab|zumab|ximab|tinib|ciclib|parib|rafenib|'
        r'lisib|metinib|lukast|sartan|pril|olol|statin))\b',
        r'\b(aspirin|ibuprofen|metformin|warfarin|heparin|insulin|penicillin|'
        r'amoxicillin|ciprofloxacin|doxorubicin|cisplatin|tamoxifen|paclitaxel|'
        r'docetaxel|gemcitabine|carboplatin|oxaliplatin|bevacizumab|'
        r'trastuzumab|rituximab|2-methoxyestradiol)\b',
        r'(?:treatment with|treated with|administered)\s+([A-Z][a-z]{3,})',
        r'([A-Z][a-z]+)\s+\([A-Z]{2,5}\)',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return "unknown"


def _extract_cell_type(text: str) -> str:
    patterns = [
        r'\b(HeLa|HEK[- ]?293|MCF-7|A549|Jurkat|CHO|PC-3|LNCaP|U87|T47D|MDA-MB-231)\b',
        r'\b([A-Za-z]+(?:\s+[A-Za-z]+)?\s+cell(?:s|\s+line))\b',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return "unknown"


def build_document(example: dict) -> dict:
    pubid = str(example.get("pubid", uuid4()))
    ctx   = example.get("context", {})
    ctx_parts = ctx.get("contexts", []) if isinstance(ctx, dict) else [str(ctx)]
    context_str = " ".join(ctx_parts) if isinstance(ctx_parts, list) else str(ctx_parts)

    question    = example.get("question", "")
    long_answer = example.get("long_answer", "")
    content = "\n\n".join(p for p in [context_str, f"Question: {question}", f"Answer: {long_answer}"] if p.strip())

    doc_id = str(uuid4())
    return {
        "content": content,
        "metadata": {
            "doc_id":                doc_id,
            "source_file":           f"pubmed_{pubid}.pdf",
            "file_type":             "pdf",
            "date_creation":         "2020-01-01",
            "year":                  2020,
            "compound_name":         _extract_compound(content),
            "test_type":             "efficacy",
            "cell_type":             _extract_cell_type(content),
            "domain":                "pharmaceutical",
            "chunk_index":           0,
            "parent_doc_id":         doc_id,
            "page_number":           1,
            "chunk_type":            "abstract",
            "language":              "en",
            "confidence_extraction": 0.90,
        },
    }


# ---------------------------------------------------------------------------
# File generators
# ---------------------------------------------------------------------------

def _safe(text: str) -> str:
    return text.encode("latin-1", "replace").decode("latin-1")


def write_txt(path: Path, title: str, content: str) -> None:
    path.write_text(f"{title}\n{'=' * len(title)}\n\n{content}", encoding="utf-8")


def write_docx(path: Path, title: str, content: str) -> None:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    h = doc.add_heading(title[:200], level=0)
    h.runs[0].font.size = Pt(14)
    doc.add_paragraph(content)
    doc.save(str(path))


def write_pdf(path: Path, title: str, content: str) -> None:
    from fpdf import FPDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 13)
    pdf.multi_cell(0, 8, _safe(title[:150]))
    pdf.ln(3)
    pdf.set_font("Helvetica", size=10)
    # Write content in blocks to avoid overflow
    for chunk in [content[i:i+2000] for i in range(0, min(len(content), 8000), 2000)]:
        pdf.multi_cell(0, 6, _safe(chunk))
    pdf.output(str(path))


def write_csv(path: Path, title: str, content: str) -> None:
    compound = _extract_compound(content)
    cell     = _extract_cell_type(content)
    if compound == "unknown":
        compound = "test_compound"
    if cell == "unknown":
        cell = "HEK-293"

    rows = [
        ["compound",      "concentration_uM", "cell_viability_pct", "IC50_uM", "cell_line", "assay",       "timepoint_h"],
        [compound,        "0.1",              "98.2",               "2.5",     cell,         "MTT",         "24"],
        [compound,        "0.5",              "91.4",               "2.5",     cell,         "MTT",         "24"],
        [compound,        "1.0",              "78.6",               "2.5",     cell,         "MTT",         "48"],
        [compound,        "2.5",              "52.1",               "2.5",     cell,         "MTT",         "48"],
        [compound,        "5.0",              "28.3",               "2.5",     cell,         "MTT",         "72"],
        [compound,        "10.0",             "11.7",               "2.5",     cell,         "MTT",         "72"],
        [compound,        "25.0",             "4.2",                "2.5",     cell,         "Flow Cytometry","72"],
    ]
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow([f"# {title}"])
        w.writerows(rows)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main(n_docs: int = N_DOCS, n_samples: int = N_SAMPLES):
    print(f"{'=' * 60}")
    print(f"  Chargeur dataset large — PubMedQA pqa_artificial")
    print(f"  Cible : {n_docs:,} documents  |  {n_samples} fichiers multi-formats")
    print(f"{'=' * 60}\n")

    # ------------------------------------------------------------------
    # 1. Load dataset
    # ------------------------------------------------------------------
    print(f"[1/3] Chargement de qiaojin/PubMedQA (pqa_artificial)...")
    print(f"      (premier telechargement ~500 MB, puis cache local)\n")
    dataset = load_dataset("qiaojin/PubMedQA", "pqa_artificial")
    split   = dataset["train"]
    total   = len(split)
    print(f"      {total:,} exemples disponibles dans le split 'train'")

    take = min(n_docs, total)
    print(f"      Selection de {take:,} documents...\n")
    examples = list(split.select(range(take)))

    # ------------------------------------------------------------------
    # 2. Build documents.json
    # ------------------------------------------------------------------
    print(f"[2/3] Construction de documents.json ({take:,} docs)...")
    documents = []
    for ex in tqdm(examples, unit="doc", desc="  Parsing"):
        documents.append(build_document(ex))

    with open(OUTPUT, "w", encoding="utf-8") as f:
        json.dump(documents, f, ensure_ascii=False, indent=2)
    print(f"      Sauvegarde -> {OUTPUT.relative_to(ROOT)}\n")

    # ------------------------------------------------------------------
    # 3. Generate multi-format sample files
    # ------------------------------------------------------------------
    print(f"[3/3] Generation de {n_samples} fichiers multi-formats dans data/raw/...")
    RAW_DIR.mkdir(parents=True, exist_ok=True)

    counts    = {k: 0 for k in FORMAT_SPLIT}
    targets   = {k: max(1, round(v * n_samples / 100)) for k, v in FORMAT_SPLIT.items()}
    writers   = {"txt": write_txt, "docx": write_docx, "pdf": write_pdf, "csv": write_csv}
    generated = 0

    for doc in tqdm(documents[:n_samples * 3], unit="file", desc="  Generating"):
        fmt = next((k for k, v in targets.items() if counts[k] < v), None)
        if fmt is None:
            break
        content = doc["content"]
        source  = doc["metadata"]["source_file"].replace(".pdf", "")
        title   = f"{source} — {doc['metadata']['compound_name']}"
        stem    = f"{source}_{fmt}_{counts[fmt]:03d}"
        path    = RAW_DIR / f"{stem}.{fmt}"
        try:
            writers[fmt](path, title, content)
            counts[fmt] += 1
            generated   += 1
        except Exception as exc:
            tqdm.write(f"  [SKIP {fmt}] {exc.__class__.__name__}: {exc}")

    print(f"\n  Fichiers generes dans data/raw/ :")
    for fmt, cnt in counts.items():
        print(f"    .{fmt:<5}  {cnt:>3} fichiers")
    print(f"  Total : {generated} fichiers\n")

    print(f"{'=' * 60}")
    print(f"  TERMINE")
    print(f"  {take:,} documents -> data/documents.json")
    print(f"  {generated} fichiers multi-formats -> data/raw/")
    print(f"{'=' * 60}")
    print(f"\nEtapes suivantes :")
    print(f"  1. Re-indexer Qdrant  : python ingestion/indexer.py")
    print(f"  2. Tester le parser   : python ingestion/parser.py")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--docs",    type=int, default=N_DOCS,    help="Nombre de documents a charger")
    parser.add_argument("--samples", type=int, default=N_SAMPLES, help="Nombre de fichiers a generer")
    args = parser.parse_args()
    main(args.docs, args.samples)
